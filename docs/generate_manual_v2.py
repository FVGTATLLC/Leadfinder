#!/usr/bin/env python3
"""Generate SalesPilot Nigeria User Manual v2 as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Global styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Ensure List Bullet style exists
try:
    lb = doc.styles['List Bullet']
except KeyError:
    lb = doc.styles.add_style('List Bullet', 1)
lb.font.name = 'Calibri'
lb.font.size = Pt(11)

DARK_BLUE = RGBColor(0x1B, 0x4F, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = "F0F0F0"
ALT_ROW = "EBF5FB"

# ─── Helper functions ───────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_screenshot_box(title, description):
    """Add a bordered gray box simulating a screenshot placeholder."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    # Border
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="8" w:space="0" w:color="1B4F72"/>'
        '  <w:left w:val="single" w:sz="8" w:space="0" w:color="1B4F72"/>'
        '  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1B4F72"/>'
        '  <w:right w:val="single" w:sz="8" w:space="0" w:color="1B4F72"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"\U0001F4F8 {title}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(description)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2.italic = True
    doc.add_paragraph()  # spacer

def add_styled_table(headers, rows, col_widths=None):
    """Add a table with dark-blue header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B4F72")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_cell_shading(cell, ALT_ROW)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_para(text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    if size:
        r.font.size = Pt(size)
    return p

def add_note_box(text):
    """Add a light-blue info box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "D6EAF8")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="16" w:space="0" w:color="2980B9"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    r = p.add_run("\u2139\uFE0F  " + text)
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    doc.add_paragraph()

# ─── COVER PAGE ─────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SalesPilot")
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = DARK_BLUE
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Nigeria Edition")
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
r.font.name = 'Calibri'

doc.add_paragraph()

# Decorative line
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.rows[0].cells[0]
set_cell_shading(cell, "1B4F72")
cell.paragraphs[0].add_run(" ")
cell.width = Inches(4)
# Make it thin
for row in table.rows:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="40" w:hRule="exact"/>')
    trPr.append(trHeight)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("User Manual")
r.font.size = Pt(28)
r.font.color.rgb = DARK_BLUE
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Outbound Sales Platform\nfor Nigerian Corporate Travel & MICE")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Club Concierge International")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("April 2026  |  Confidential")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Page break after cover
doc.add_page_break()

# ─── HEADERS & FOOTERS (from section 1 onward) ─────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.different_first_page_header_footer = True

# Regular header
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("SalesPilot \u2014 User Manual")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.name = 'Calibri'
hp.add_run("    ")
r2 = hp.add_run("Confidential")
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
r2.bold = True
r2.font.name = 'Calibri'

# Footer with page number
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Page ")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
# Add page number field
fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_elem = fp.add_run()._r
run_elem.append(fld_char1)
run_elem2 = fp.add_run()._r
run_elem2.append(instr)
run_elem3 = fp.add_run()._r
run_elem3.append(fld_char2)

# ─── TABLE OF CONTENTS ─────────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph()

toc_items = [
    ("1", "Introduction", ""),
    ("2", "Getting Started", ""),
    ("3", "Dashboard", ""),
    ("4", "Strategies (ICP)", ""),
    ("5", "Companies", ""),
    ("6", "Contacts", ""),
    ("7", "Campaigns", ""),
    ("8", "Messages", ""),
    ("9", "Research", ""),
    ("10", "Analytics", ""),
    ("11", "Data Exports", ""),
    ("12", "Settings", ""),
    ("13", "Admin Panel", ""),
    ("14", "Troubleshooting & Support", ""),
]

for num, title, _ in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"Chapter {num}:  {title}")
    r.font.size = Pt(12)
    r.font.name = 'Calibri'
    if num in ("1", "2", "3"):
        r.font.color.rgb = DARK_BLUE

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 1: Introduction
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 1: Introduction', level=1)

doc.add_heading('1.1  What Is SalesPilot?', level=2)
doc.add_paragraph(
    'SalesPilot is an AI-powered outbound sales management platform purpose-built for the Nigerian '
    'corporate travel and MICE (Meetings, Incentives, Conferences & Exhibitions) market. It helps '
    'sales teams identify, research, and engage high-value corporate prospects across Nigeria through '
    'intelligent automation, AI-generated messaging, and a structured sales workflow.'
)
doc.add_paragraph(
    'The platform is operated by Club Concierge International and is accessible as a web application '
    'at the following URL:'
)
add_note_box("Application URL:  https://sales-management-nigeria.vercel.app")

doc.add_heading('1.2  Key Features', level=2)
features = [
    "Ideal Customer Profile (ICP) Strategy Builder with AI-driven company discovery",
    "Nigerian Company Database with ICP scoring and enrichment",
    "AI Persona Discovery for finding key contacts with Nigerian naming conventions",
    "Multi-step Campaign Builder with customizable tones and templates",
    "AI Message Generation with per-contact personalization",
    "Research Briefs powered by AI for company intelligence and talking points",
    "Analytics Dashboard with funnel visualization, leaderboards, and trend charts",
    "Admin Panel for user management, approvals, and activity auditing",
]
for f in features:
    add_bullet(f)

doc.add_heading('1.3  User Roles', level=2)
doc.add_paragraph(
    'SalesPilot supports four user roles, each with different levels of access to the platform. '
    'Roles are assigned during registration approval by an administrator.'
)
add_styled_table(
    ["Role", "Access Level", "Description"],
    [
        ["Admin", "Full", "Complete access to all features including the Admin Panel, user management, approvals, and system settings."],
        ["Manager", "High", "Access to all sales features, analytics, team performance data, and campaign approval rights."],
        ["Sales Rep", "Standard", "Access to companies, contacts, campaigns, messages, and research. Cannot approve users or access admin settings."],
        ["Viewer", "Read-Only", "Can view dashboards, companies, and reports but cannot create or modify records."],
    ],
    col_widths=[1.2, 1.0, 4.3]
)

doc.add_heading('1.4  System Requirements', level=2)
add_bullet("Modern web browser (Chrome 90+, Firefox 88+, Safari 14+, Edge 90+)")
add_bullet("Stable internet connection")
add_bullet("Screen resolution of 1280x720 or higher recommended")
add_bullet("No software installation required \u2014 SalesPilot is fully web-based")

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 2: Getting Started
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 2: Getting Started', level=1)

doc.add_heading('2.1  Logging In', level=2)
doc.add_paragraph(
    'To access SalesPilot, navigate to the application URL in your browser. You will be presented '
    'with the login screen.'
)

add_screenshot_box(
    "[Screenshot: Login Page]",
    "The login screen displays the SalesPilot logo at the top, a \"Welcome back\" heading, "
    "a Login ID field with placeholder text \"you@clubconcierge.com\", a Password field with "
    "visibility toggle, a blue \"Sign in\" button, and a \"Don't have an account? Create one\" "
    "link at the bottom for new user registration."
)

doc.add_paragraph('To log in:')
add_bullet("Enter your Login ID (email address) in the Login ID field.")
add_bullet("Enter your password in the Password field.")
add_bullet("Click the \"Sign in\" button.")
add_bullet("If your credentials are valid, you will be redirected to the Dashboard.")

add_note_box(
    "First-time login: After your account is approved by an administrator, you will receive "
    "a generated password. You will be required to change this password on your first login."
)

doc.add_heading('2.2  Registering a New Account', level=2)
doc.add_paragraph(
    'If you do not yet have an account, click the "Create one" link on the login page to access '
    'the registration form.'
)

add_screenshot_box(
    "[Screenshot: Registration Page]",
    "The registration screen shows a \"Request an account\" heading with an informational box "
    "stating \"Your account request will be reviewed by an administrator. You'll receive your login "
    "credentials once approved.\" Below this are fields for Full Name, Email Address (placeholder: "
    "you@clubconcierge.com), a Role dropdown defaulting to \"Sales Representative\", and a blue "
    "\"Submit Registration\" button."
)

doc.add_paragraph('The registration-approval flow works as follows:')
add_bullet("Step 1: Fill in your Full Name, Email Address, and select your desired Role.")
add_bullet("Step 2: Click \"Submit Registration\" to send your request.")
add_bullet("Step 3: An administrator reviews your request in the Admin Panel.")
add_bullet("Step 4: Upon approval, the system generates a temporary password.")
add_bullet("Step 5: You receive your login credentials (typically via your manager or email).")
add_bullet("Step 6: Log in with the temporary password and set a new password on first login.")

doc.add_heading('2.3  Navigation Sidebar', level=2)
doc.add_paragraph(
    'Once logged in, the left sidebar provides access to all major sections of the application. '
    'The sidebar is collapsible and highlights the currently active section.'
)

add_styled_table(
    ["Menu Item", "Icon", "Description"],
    [
        ["Dashboard", "Home", "Overview of KPIs, sales funnel, and recent campaigns"],
        ["Strategies", "Target", "ICP strategy management and company discovery"],
        ["Companies", "Building", "Company database with ICP scoring"],
        ["Contacts", "Users", "Contact management with AI persona discovery"],
        ["Campaigns", "Megaphone", "Multi-step campaign builder"],
        ["Messages", "Mail", "Message queue with approval workflow"],
        ["Research", "Microscope", "AI-generated research briefs"],
        ["Analytics", "Chart", "Performance dashboards and trend analysis"],
        ["Exports", "Download", "Data export functionality"],
        ["Settings", "Gear", "Application and integration settings"],
        ["Admin", "Shield", "User management and activity logs (Admin only)"],
    ],
    col_widths=[1.4, 0.8, 4.3]
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 3: Dashboard
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 3: Dashboard', level=1)

doc.add_paragraph(
    'The Dashboard is the landing page after login. It provides a high-level overview of your '
    'sales pipeline, key performance indicators, and recent activity.'
)

add_screenshot_box(
    "[Screenshot: Dashboard]",
    "The dashboard displays a \"Welcome back, Admin\" greeting with the current date. Below are "
    "6 KPI cards arranged in a grid: Strategies (3), Companies (142), Contacts (21), Active "
    "Campaigns (0), Messages Sent (0), and Response Rate (0.0%). A Sales Funnel section shows "
    "the pipeline flow from Strategies to Companies to Contacts. The Recent Campaigns section "
    "lists \"test\" and \"Apr Campaign\" both marked as Draft status with gray badges."
)

doc.add_heading('3.1  KPI Cards', level=2)
doc.add_paragraph(
    'The six KPI cards at the top of the dashboard provide real-time metrics for your sales operation:'
)

add_styled_table(
    ["KPI Card", "Description", "Example Value"],
    [
        ["Strategies", "Total number of ICP strategies created", "3"],
        ["Companies", "Total companies in the database across all strategies", "142"],
        ["Contacts", "Total contacts added or discovered", "21"],
        ["Active Campaigns", "Campaigns currently in Active status", "0"],
        ["Messages Sent", "Total messages sent across all campaigns", "0"],
        ["Response Rate", "Percentage of sent messages that received replies", "0.0%"],
    ],
    col_widths=[1.5, 3.5, 1.2]
)

doc.add_heading('3.2  Sales Funnel', level=2)
doc.add_paragraph(
    'The Sales Funnel section provides a visual representation of your outbound pipeline, showing '
    'the progression from broad ICP strategies down to individual contacts. Each stage shows the '
    'count and the conversion percentage between stages. This helps you quickly identify where '
    'your pipeline may need attention.'
)

doc.add_heading('3.3  Recent Campaigns', level=2)
doc.add_paragraph(
    'The Recent Campaigns section displays the most recently created or modified campaigns, '
    'including their name, status badge (Draft, Active, Paused, or Completed), and key statistics '
    'such as the number of contacts and message steps. Click on any campaign to navigate directly '
    'to its detail page.'
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 4: Strategies (ICP)
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 4: Strategies (ICP)', level=1)

doc.add_paragraph(
    'The Strategies module is where you define your Ideal Customer Profile (ICP) and use it to '
    'discover and organize target companies. Each strategy represents a specific market segment '
    'you want to pursue.'
)

add_screenshot_box(
    "[Screenshot: ICP Strategies Page]",
    "The page displays an \"ICP Strategies\" heading with a blue \"Create Strategy\" button in the "
    "top right. Below are 3 strategy cards: \"Fintech Companies\" (Active status with green badge, "
    "29 companies), \"FMCG Companies in lagos\" (Active, 20 companies), and \"Top 20\" (Draft "
    "status with gray badge, 10 companies). Each card shows the strategy name, status, company "
    "count, and a brief description."
)

doc.add_heading('4.1  Creating a Strategy', level=2)
doc.add_paragraph('To create a new ICP strategy:')
add_bullet("Click the \"Create Strategy\" button in the top-right corner.")
add_bullet("Enter a Strategy Name (e.g., \"Fintech Companies in Lagos\").")
add_bullet("Provide a Description explaining the target market segment.")
add_bullet("Configure ICP Filters to define the characteristics of your ideal companies.")
add_bullet("Click \"Save\" to create the strategy in Draft status.")

doc.add_heading('4.2  ICP Filters', level=2)
doc.add_paragraph(
    'ICP filters allow you to define the characteristics of companies that match your ideal '
    'customer profile. The following filters are available:'
)

add_styled_table(
    ["Filter", "Type", "Description", "Example Values"],
    [
        ["Industry", "Multi-select", "Target industry verticals", "Fintech, Oil & Gas, Banking"],
        ["City", "Multi-select", "Geographic location within Nigeria", "Lagos, Abuja, Port Harcourt"],
        ["Employee Range", "Range", "Company size by headcount", "50\u2013500, 500\u20135000"],
        ["Revenue Range", "Range", "Estimated annual revenue", "\u20A61B\u2013\u20A610B"],
        ["Travel Intensity", "Select", "Estimated corporate travel volume", "High, Medium, Low"],
        ["Custom Tags", "Free text", "Additional tags for segmentation", "Multinational, Listed, SME"],
    ],
    col_widths=[1.3, 1.0, 2.5, 1.8]
)

doc.add_heading('4.3  AI Company Discovery', level=2)
doc.add_paragraph(
    'Once your ICP filters are configured, SalesPilot can use AI to automatically discover '
    'Nigerian companies that match your criteria. The AI searches across multiple data sources '
    'and returns companies with their ICP match score, industry classification, and estimated '
    'size. Discovered companies are added to the strategy and become available in the Companies '
    'module for further enrichment and outreach.'
)

add_note_box(
    "AI Discovery uses web intelligence and business databases to find companies. Results are "
    "specific to the Nigerian market and prioritize companies with corporate travel potential."
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 5: Companies
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 5: Companies', level=1)

doc.add_paragraph(
    'The Companies module is the central database of all prospect companies in SalesPilot. '
    'Companies can be added through AI discovery, manual entry, or CSV import, and each '
    'receives an ICP score based on how well it matches your strategy criteria.'
)

add_screenshot_box(
    "[Screenshot: Companies Page]",
    "A data table with columns: Company, Industry, Geography, Employees, ICP Score, Source, "
    "and Added date. Sample rows include: Verite Microfinance Bank (Financial Services, Lagos, "
    "200 employees), Africa Prudential Plc (Financial Services, Lagos, 300), Payant (Fintech, "
    "Lagos, 150), Wema Bank Plc (Banking & Finance, Lagos, 3000), Union Bank of Nigeria Plc "
    "(Banking & Finance, Lagos, 3500). Above the table: a search bar, Filters dropdown button, "
    "\"Import CSV\" button, and blue \"Add Company\" button."
)

doc.add_heading('5.1  Adding Companies', level=2)
doc.add_paragraph('There are three methods to add companies to SalesPilot:')

p = doc.add_paragraph()
r = p.add_run('Method 1: AI Discovery')
r.bold = True
doc.add_paragraph(
    'Use the AI Company Discovery feature within a Strategy to automatically find and add '
    'companies matching your ICP criteria. This is the recommended approach for building your '
    'initial prospect list.'
)

p = doc.add_paragraph()
r = p.add_run('Method 2: Manual Entry')
r.bold = True
doc.add_paragraph(
    'Click the "Add Company" button to open the manual entry form. Fill in the company details '
    'and click Save. This is useful for adding specific companies you already know about.'
)

p = doc.add_paragraph()
r = p.add_run('Method 3: CSV Import')
r.bold = True
doc.add_paragraph(
    'Click the "Import CSV" button to upload a spreadsheet of companies. The CSV file should '
    'include columns for company name, industry, city, employee count, and other relevant fields. '
    'SalesPilot will map the columns and import the data.'
)

doc.add_heading('5.2  Company Fields', level=2)
add_styled_table(
    ["Field", "Required", "Description"],
    [
        ["Company Name", "Yes", "Official registered name of the company"],
        ["Industry", "Yes", "Industry vertical (e.g., Fintech, Oil & Gas, Banking)"],
        ["City / Geography", "Yes", "Primary office location in Nigeria"],
        ["Employee Count", "No", "Estimated number of employees"],
        ["Revenue Estimate", "No", "Estimated annual revenue in Naira"],
        ["Website", "No", "Company website URL"],
        ["Description", "No", "Brief description of the company's business"],
        ["Travel Intensity", "No", "Estimated corporate travel volume (High/Medium/Low)"],
        ["Custom Tags", "No", "User-defined tags for segmentation and filtering"],
        ["Source", "Auto", "How the company was added (AI, Manual, CSV)"],
    ],
    col_widths=[1.5, 0.8, 4.2]
)

doc.add_heading('5.3  ICP Score', level=2)
doc.add_paragraph(
    'Every company receives an ICP Score from 0 to 100 that indicates how well it matches your '
    'Ideal Customer Profile. The score is calculated using a weighted formula:'
)

add_styled_table(
    ["Factor", "Weight", "Description"],
    [
        ["Industry Match", "30%", "Whether the company's industry aligns with your target verticals"],
        ["Company Size", "20%", "How the employee count fits your target range"],
        ["Revenue Fit", "20%", "Whether estimated revenue matches your target range"],
        ["City Match", "15%", "Whether the company is located in one of your target cities"],
        ["Travel Intensity", "15%", "Estimated corporate travel spend and frequency"],
    ],
    col_widths=[1.5, 0.8, 4.2]
)

doc.add_heading('5.4  Company Detail Page', level=2)
doc.add_paragraph(
    'Clicking on any company row opens the Company Detail page, which provides a comprehensive '
    'view of all information about that company, including:'
)
add_bullet("Company overview with all fields and ICP score breakdown")
add_bullet("Associated contacts with their roles and enrichment status")
add_bullet("Campaign history showing all campaigns the company has been included in")
add_bullet("AI-generated research briefs and talking points")
add_bullet("Activity timeline with all interactions and status changes")

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 6: Contacts
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 6: Contacts', level=1)

doc.add_paragraph(
    'The Contacts module manages individual prospects within target companies. Contacts can be '
    'added manually or discovered using AI Persona Discovery, which finds key decision-makers '
    'using Nigerian naming conventions and business intelligence.'
)

add_screenshot_box(
    "[Screenshot: Contacts Page]",
    "A data table with columns: Name, Email, Job Title, Persona, Company, Confidence, "
    "Enrichment, and LinkedIn. Sample contacts include: Adekunle Adeyemi (CFO at Opay), "
    "Chinyere Okoro (HR Head at Chevron Nigeria Limited), and Chinyere Okoro (HR Head at "
    "Cappa & D'Alberto Plc). The Persona column shows colored badges (e.g., Finance Leader "
    "in blue, HR Leader in green). The LinkedIn column displays clickable profile icons. "
    "Above the table: search bar, Filters dropdown, and \"Add Contact\" button."
)

doc.add_heading('6.1  Adding Contacts', level=2)

p = doc.add_paragraph()
r = p.add_run('Manual Entry')
r.bold = True
doc.add_paragraph(
    'Click "Add Contact" to open the contact form. Enter the contact\'s name, email, job title, '
    'company association, and other details. The contact will be added and associated with the '
    'selected company.'
)

p = doc.add_paragraph()
r = p.add_run('AI Persona Discovery')
r.bold = True
doc.add_paragraph(
    'AI Persona Discovery is an intelligent feature that automatically finds potential contacts '
    'at target companies. The AI uses Nigerian naming conventions, business directories, and '
    'web intelligence to discover key decision-makers. Each discovered contact includes a '
    'confidence score indicating the reliability of the information found.'
)

doc.add_heading('6.2  Contact Fields', level=2)
add_styled_table(
    ["Field", "Required", "Description"],
    [
        ["Full Name", "Yes", "Contact's full name"],
        ["Email", "Yes", "Business email address"],
        ["Job Title", "Yes", "Current position at the company"],
        ["Persona", "Auto", "AI-assigned persona category (e.g., Finance Leader, HR Leader, C-Suite)"],
        ["Company", "Yes", "Associated company from the Companies database"],
        ["Phone", "No", "Direct phone number"],
        ["LinkedIn URL", "No", "LinkedIn profile URL for research and outreach"],
        ["Confidence", "Auto", "AI confidence score for discovered contacts (0\u2013100%)"],
        ["Enrichment Status", "Auto", "Whether the contact has been enriched with additional data"],
    ],
    col_widths=[1.5, 0.8, 4.2]
)

doc.add_heading('6.3  Contact Enrichment', level=2)
doc.add_paragraph(
    'Contact enrichment adds additional data points to a contact record, including verified email '
    'addresses, phone numbers, social media profiles, and recent professional activity. Enrichment '
    'can be triggered manually for individual contacts or in bulk for an entire strategy.'
)

doc.add_heading('6.4  AI Persona Discovery', level=2)
doc.add_paragraph(
    'The AI Persona Discovery feature is specifically designed for the Nigerian market. It:'
)
add_bullet("Recognizes Nigerian naming conventions (Yoruba, Igbo, Hausa, and other ethnic names)")
add_bullet("Identifies key decision-makers by title and seniority level")
add_bullet("Assigns persona categories (Finance Leader, HR Leader, Travel Manager, C-Suite, Procurement)")
add_bullet("Provides a confidence score based on the quality and consistency of source data")
add_bullet("Links to LinkedIn profiles where available for further verification")

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 7: Campaigns
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 7: Campaigns', level=1)

doc.add_paragraph(
    'Campaigns are the core outreach mechanism in SalesPilot. A campaign brings together a '
    'strategy, a set of contacts, and a sequence of personalized messages to execute a '
    'structured outbound sales motion.'
)

add_screenshot_box(
    "[Screenshot: Campaigns Page]",
    "The campaigns page shows 2 campaign cards: \"test\" (Corporate type, Draft status, "
    "Consultative tone, 0 contacts, 0 steps) and \"Apr Campaign\" (Intro type, Draft status, "
    "Consultative tone, 0 contacts, 0 steps). Above the cards: a search bar, Status filter "
    "dropdown, Type filter dropdown, and a blue \"Create Campaign\" button in the top right."
)

doc.add_heading('7.1  Creating a Campaign', level=2)
doc.add_paragraph('To create a new campaign:')
add_bullet("Click the \"Create Campaign\" button.")
add_bullet("Enter a Campaign Name that describes the outreach objective.")
add_bullet("Select a Strategy to determine which companies and contacts are targeted.")
add_bullet("Choose a Campaign Type (see types below).")
add_bullet("Select a Tone Preset for the messaging style.")
add_bullet("Save the campaign. It will be created in Draft status.")

doc.add_heading('7.2  Campaign Types', level=2)
add_styled_table(
    ["Type", "Purpose", "Typical Use Case"],
    [
        ["Introduction", "First outreach to new prospects", "Initial contact with companies from a new ICP strategy"],
        ["Follow-Up", "Re-engage previous contacts", "Second or third touch after an introduction campaign"],
        ["MICE", "Meetings & events focused", "Promoting conference, incentive travel, or event services"],
        ["Corporate", "Corporate travel services", "Selling managed corporate travel programs"],
        ["Custom", "User-defined campaign type", "Any outreach that doesn't fit the above categories"],
    ],
    col_widths=[1.2, 2.2, 3.2]
)

doc.add_heading('7.3  Tone Presets', level=2)
add_styled_table(
    ["Tone", "Style", "Best For"],
    [
        ["Formal", "Professional, structured language", "C-suite executives, banking, government"],
        ["Friendly", "Warm, conversational, approachable", "HR leaders, SME owners, repeat contacts"],
        ["Consultative", "Advisory, insight-led, value-focused", "Finance leaders, procurement heads"],
        ["Assertive", "Direct, confident, action-oriented", "Follow-ups, time-sensitive offers"],
    ],
    col_widths=[1.2, 2.5, 2.8]
)

doc.add_heading('7.4  Campaign Sequence Steps', level=2)
doc.add_paragraph(
    'Each campaign contains one or more sequence steps that define the outreach cadence. Steps '
    'can include initial emails, follow-ups, and LinkedIn connection requests. Each step has a '
    'delay period (e.g., wait 3 days) before execution.'
)

doc.add_heading('7.5  Campaign Statuses', level=2)
add_styled_table(
    ["Status", "Badge Color", "Description"],
    [
        ["Draft", "Gray", "Campaign is being configured. Not yet active."],
        ["Active", "Green", "Campaign is live and messages are being sent according to the sequence."],
        ["Paused", "Yellow", "Campaign is temporarily suspended. Can be resumed."],
        ["Completed", "Blue", "All sequence steps have been executed for all contacts."],
    ],
    col_widths=[1.2, 1.2, 4.2]
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 8: Messages
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 8: Messages', level=1)

doc.add_paragraph(
    'The Messages module is the central hub for all outbound communications generated by '
    'campaigns. Messages go through an approval workflow before being sent, ensuring quality '
    'control over all prospect communications.'
)

add_screenshot_box(
    "[Screenshot: Messages Page]",
    "The messages page displays a \"Messages\" heading with filter tabs: All, Drafts, Pending "
    "Approval, Approved, Sent, and Failed. Below the tabs are a search bar and a campaign filter "
    "dropdown. The main content area shows an empty state message: \"No messages yet \u2014 Messages "
    "will appear here once you generate them from campaigns.\" The empty state includes a mail "
    "icon illustration."
)

doc.add_heading('8.1  Message Generation', level=2)
doc.add_paragraph(
    'Messages are generated by AI when a campaign is activated. The AI uses the campaign type, '
    'tone preset, contact information, company data, and research briefs to craft personalized '
    'messages for each contact. Each message is tailored to the specific prospect, referencing '
    'their company, role, and potential pain points.'
)

doc.add_heading('8.2  Approval Workflow', level=2)
doc.add_paragraph(
    'All AI-generated messages follow a structured approval workflow to ensure quality and '
    'appropriateness before reaching prospects:'
)

add_styled_table(
    ["Status", "Description", "Next Action"],
    [
        ["Draft", "Message has been generated by AI", "Review and submit for approval"],
        ["Pending Approval", "Message awaits manager/admin review", "Approve or reject"],
        ["Approved", "Message has been approved for sending", "Schedule or send immediately"],
        ["Scheduled", "Message is queued for delivery at a future date/time", "Automatic sending"],
        ["Sent", "Message has been delivered to the prospect", "Monitor for response"],
        ["Replied", "Prospect has responded to the message", "Follow up manually"],
        ["Bounced", "Message could not be delivered (invalid email)", "Verify contact email"],
        ["Failed", "Message delivery failed (technical error)", "Retry or investigate"],
    ],
    col_widths=[1.3, 3.0, 2.2]
)

doc.add_heading('8.3  Message Actions', level=2)
add_styled_table(
    ["Action", "Description"],
    [
        ["Preview", "View the full message content in a modal window"],
        ["Edit", "Modify the AI-generated message text before sending"],
        ["Regenerate", "Ask the AI to create a new version of the message"],
        ["Approve", "Mark the message as approved for delivery (Manager/Admin)"],
        ["Reject", "Return the message to draft with feedback for revision"],
        ["Schedule", "Set a specific date and time for message delivery"],
        ["Bulk Actions", "Select multiple messages to approve, reject, or schedule at once"],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 9: Research
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 9: Research', level=1)

doc.add_paragraph(
    'The Research module provides AI-generated intelligence briefs about companies and prospects. '
    'These briefs equip sales representatives with the context and talking points needed for '
    'effective outreach conversations.'
)

add_screenshot_box(
    "[Screenshot: Research Briefs Page]",
    "The page shows a \"Research Briefs\" heading with a blue \"Generate New Research\" button. "
    "Below are 6 research brief cards arranged in a grid: Medbury (Company Summary, generated "
    "by Agent), Club Concierge International (Company Summary), Baker Hughes Nigeria (Company "
    "Summary), Paystack (2 separate brief cards), and PwC Nigeria. Each card shows the company "
    "name, brief type label, a summary paragraph, and key facts section."
)

doc.add_heading('9.1  Brief Types', level=2)
add_styled_table(
    ["Brief Type", "Content Focus", "Best Used For"],
    [
        ["Company Summary", "Overview, history, key facts, and market position", "Initial research before first outreach"],
        ["Prospect Summary", "Individual prospect's background and role context", "Personalizing messages to a specific contact"],
        ["Talking Points", "Conversation starters and value propositions", "Preparing for calls or meetings"],
        ["Industry Brief", "Sector-level trends and challenges", "Understanding the prospect's market context"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading('9.2  Generating a Research Brief', level=2)
doc.add_paragraph('To generate a new research brief:')
add_bullet("Click the \"Generate New Research\" button.")
add_bullet("Select the target company from the dropdown.")
add_bullet("Choose the brief type (Company Summary, Talking Points, etc.).")
add_bullet("Optionally select a specific contact for prospect-focused briefs.")
add_bullet("Click \"Generate\" and wait for the AI to compile the research.")

doc.add_heading('9.3  Research Output Structure', level=2)
doc.add_paragraph('Each research brief contains the following sections:')

add_styled_table(
    ["Section", "Description"],
    [
        ["Summary", "A concise 2\u20133 paragraph overview of the company or prospect"],
        ["Key Facts", "Bullet-point list of important data (industry, size, revenue, location)"],
        ["Talking Points", "Suggested conversation topics relevant to the prospect's role"],
        ["Pain Points", "Identified business challenges the prospect likely faces"],
        ["Opportunities", "Potential areas where Club Concierge services can add value"],
        ["Recent News", "Latest news mentions or developments about the company"],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_heading('9.4  Nigeria-Specific Intelligence', level=2)
doc.add_paragraph(
    'Research briefs generated by SalesPilot include intelligence specific to the Nigerian '
    'corporate travel and MICE market. Common pain points identified include:'
)
add_bullet("FX volatility and its impact on international travel budgets")
add_bullet("Duty of care obligations for employees traveling domestically and internationally")
add_bullet("Visa processing challenges for Nigerian business travelers")
add_bullet("Travel policy compliance and cost control in a high-inflation environment")
add_bullet("Local airline reliability and alternative routing for domestic travel")
add_bullet("Security considerations for travel to certain Nigerian regions")

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 10: Analytics
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 10: Analytics', level=1)

doc.add_paragraph(
    'The Analytics module provides comprehensive dashboards and visualizations to track sales '
    'performance, pipeline health, team productivity, and outreach effectiveness.'
)

add_screenshot_box(
    "[Screenshot: Analytics Page]",
    "The analytics page shows KPI cards at the top: Strategies (3), Companies (142), Contacts "
    "(21), Active Campaigns (0), Messages Sent (0), Response Rate (0.0%). Below is a Sales Funnel "
    "visualization: Strategies 3 \u2192 Companies 142 (+4733.3%) \u2192 Contacts 21 (+14.8%) \u2192 Enriched 0. "
    "A Rep Leaderboard section (currently empty) and a Trends section with a \"Messages Sent\" "
    "dropdown and Daily/Weekly/Monthly toggle buttons appear below the funnel."
)

doc.add_heading('10.1  Available Views', level=2)
add_styled_table(
    ["View", "Description"],
    [
        ["Sales Funnel", "Visual pipeline from strategies through companies, contacts, to enriched prospects with conversion percentages"],
        ["Campaign Performance", "Per-campaign metrics including open rates, reply rates, and bounce rates"],
        ["Rep Leaderboard", "Ranked list of sales representatives by key performance metrics"],
        ["Trends", "Time-series charts for messages sent, response rates, and pipeline growth with daily/weekly/monthly granularity"],
        ["AI Insights", "AI-generated observations about pipeline trends and recommended actions"],
    ],
    col_widths=[1.8, 4.8]
)

doc.add_heading('10.2  Key Metrics', level=2)
add_styled_table(
    ["Metric", "Calculation", "Target"],
    [
        ["Pipeline Conversion", "Contacts \u00f7 Companies", "> 15%"],
        ["Enrichment Rate", "Enriched Contacts \u00f7 Total Contacts", "> 80%"],
        ["Message Delivery Rate", "Sent \u00f7 (Sent + Bounced + Failed)", "> 95%"],
        ["Response Rate", "Replied \u00f7 Sent", "> 5%"],
        ["Campaign Completion Rate", "Completed Campaigns \u00f7 Total Campaigns", "> 70%"],
    ],
    col_widths=[1.8, 2.8, 1.0]
)

doc.add_heading('10.3  AI Insights', level=2)
doc.add_paragraph(
    'The AI Insights feature analyzes your pipeline data and provides actionable recommendations. '
    'Examples of insights include identifying underperforming campaigns, suggesting optimal send '
    'times based on response patterns, and highlighting strategies that are producing the highest '
    'quality leads.'
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 11: Data Exports
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 11: Data Exports', level=1)

doc.add_paragraph(
    'The Data Exports module allows you to export data from SalesPilot in various formats '
    'for external analysis, reporting, or integration with other systems.'
)

add_screenshot_box(
    "[Screenshot: Data Exports Page]",
    "The page shows a \"Data Exports\" heading with a blue \"New Export\" button. The main area "
    "displays an empty state message: \"No exports yet \u2014 Generate your first data export to "
    "download companies, contacts, or campaign reports.\" An illustration of a download icon "
    "accompanies the empty state text."
)

doc.add_heading('11.1  Export Types', level=2)
add_styled_table(
    ["Export Type", "Contents", "Format"],
    [
        ["Companies", "All company records with ICP scores, industry, size, and tags", "CSV / XLSX"],
        ["Contacts", "All contact records with email, title, persona, and enrichment data", "CSV / XLSX"],
        ["Activities", "Timeline of all activities, status changes, and interactions", "CSV / XLSX"],
        ["Full CRM", "Complete database export including companies, contacts, and activities", "XLSX"],
        ["Campaign Report", "Per-campaign performance metrics and message delivery status", "CSV / XLSX / PDF"],
    ],
    col_widths=[1.5, 3.2, 1.3]
)

doc.add_heading('11.2  How to Export', level=2)
add_bullet("Click the \"New Export\" button.")
add_bullet("Select the export type from the dropdown menu.")
add_bullet("Optionally apply filters (e.g., date range, strategy, campaign).")
add_bullet("Choose the output format (CSV, XLSX, or PDF where available).")
add_bullet("Click \"Generate Export\" and wait for processing.")
add_bullet("Once ready, click \"Download\" to save the file to your computer.")

add_note_box(
    "Large exports may take a few moments to generate. You will receive a notification "
    "when your export is ready for download."
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 12: Settings
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 12: Settings', level=1)

doc.add_paragraph(
    'The Settings page allows you to configure application preferences, manage integrations, '
    'and update your profile.'
)

add_screenshot_box(
    "[Screenshot: Settings Page]",
    "The Settings page displays tabs at the top: General, Team, Integrations, and Profile. "
    "The Integrations tab is active, showing a \"Zoho Mail Connection\" section with a status "
    "indicator reading \"Zoho Mail not connected\" and a blue \"Connect Zoho Mail\" button. "
    "Below is an info box explaining the SMTP connection via smtp.zoho.com. Further down, an "
    "\"App Configuration\" section shows Default Tone Preset set to \"Formal\" and Default "
    "Campaign Type set to \"Introduction\" in dropdown selectors."
)

doc.add_heading('12.1  Zoho Mail Connection', level=2)
doc.add_paragraph(
    'SalesPilot uses Zoho Mail for sending outbound emails. To connect your Zoho Mail account:'
)
add_bullet("Navigate to Settings > Integrations.")
add_bullet("Click the \"Connect Zoho Mail\" button.")
add_bullet("Enter your Zoho Mail email address and app-specific password.")
add_bullet("SalesPilot will verify the connection using smtp.zoho.com.")
add_bullet("Once connected, the status indicator will change to \"Connected\" with a green badge.")

add_note_box(
    "You must generate an app-specific password in your Zoho Mail account settings. "
    "Do not use your regular Zoho login password. Go to Zoho Accounts > Security > "
    "App Passwords to generate one."
)

doc.add_heading('12.2  App Configuration', level=2)
add_styled_table(
    ["Setting", "Options", "Description"],
    [
        ["Default Tone Preset", "Formal, Friendly, Consultative, Assertive", "Sets the default tone for new campaigns"],
        ["Default Campaign Type", "Introduction, Follow-Up, MICE, Corporate, Custom", "Sets the default type when creating new campaigns"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_heading('12.3  Profile Settings', level=2)
doc.add_paragraph(
    'The Profile tab allows you to update your personal information and change your password:'
)
add_bullet("Full Name: Update your display name across the application.")
add_bullet("Email Address: View your registered email (cannot be changed).")
add_bullet("Change Password: Enter your current password and set a new one.")
add_bullet("Notification Preferences: Configure which events trigger email or in-app notifications.")

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 13: Admin Panel
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 13: Admin Panel', level=1)

doc.add_paragraph(
    'The Admin Panel is accessible only to users with the Admin role. It provides tools for '
    'user management, registration approval, and system monitoring.'
)

add_screenshot_box(
    "[Screenshot: Admin Panel]",
    "The Admin Panel displays 4 tabs: Overview, Pending Approvals, All Users, and Activity Log. "
    "The Overview tab is active, showing 4 stats cards: Total Users (--), Active Users (--), "
    "Pending Approvals (--), and Online Now (--). Below is a \"Recent Registrations (Pending)\" "
    "section showing \"No pending registrations.\" followed by an \"Active Sessions\" table with "
    "one row: admin@clubconcierge.com with the Admin role badge."
)

doc.add_heading('13.1  User Management', level=2)
doc.add_paragraph(
    'The Admin Panel provides a complete view of all users in the system. Administrators can:'
)
add_bullet("View all registered users with their roles, status, and last login date")
add_bullet("Edit user roles (promote or demote between Admin, Manager, Sales Rep, Viewer)")
add_bullet("Activate or deactivate user accounts")
add_bullet("Reset user passwords (generates a new temporary password)")
add_bullet("Delete user accounts (soft delete with data preservation)")

doc.add_heading('13.2  Approval Workflow', level=2)
doc.add_paragraph(
    'When a new user registers through the registration form, their account enters a pending '
    'state. Administrators manage approvals through the Pending Approvals tab:'
)

p = doc.add_paragraph()
r = p.add_run('Approving a Registration:')
r.bold = True
add_bullet("Review the applicant's name, email, and requested role.")
add_bullet("Click \"Approve\" to activate the account.")
add_bullet("The system generates a temporary password and displays it.")
add_bullet("Communicate the password to the new user securely.")
add_bullet("The user will be required to change the password on first login.")

p = doc.add_paragraph()
r = p.add_run('Rejecting a Registration:')
r.bold = True
add_bullet("Click \"Reject\" to deny the registration request.")
add_bullet("The account is soft-deleted from the system.")
add_bullet("The applicant's email is freed for future registration attempts.")

doc.add_heading('13.3  Creating Users Directly', level=2)
doc.add_paragraph(
    'Administrators can also create user accounts directly without going through the registration '
    'flow. This is useful for onboarding team members quickly. Click "Add User" in the All Users '
    'tab, fill in the user details including their role, and the system will generate credentials.'
)

doc.add_heading('13.4  Activity Logs', level=2)
doc.add_paragraph(
    'The Activity Log tab provides a chronological record of all significant actions taken in '
    'the system, including:'
)
add_bullet("User logins and logouts with timestamps and IP addresses")
add_bullet("Account approvals, rejections, and role changes")
add_bullet("Campaign creation, activation, and completion events")
add_bullet("Message approvals, sends, and delivery status changes")
add_bullet("Data imports, exports, and bulk operations")
add_bullet("Settings changes and integration connection events")

doc.add_heading('13.5  Active Sessions', level=2)
doc.add_paragraph(
    'The Active Sessions section in the Overview tab shows all currently logged-in users with '
    'their email address, role, and session duration. Administrators can forcibly end a session '
    'if needed for security purposes.'
)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
# CHAPTER 14: Troubleshooting & Support
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Chapter 14: Troubleshooting & Support', level=1)

doc.add_heading('14.1  Common Problems & Solutions', level=2)

add_styled_table(
    ["Problem", "Possible Cause", "Solution"],
    [
        ["Cannot log in", "Incorrect credentials or account not yet approved",
         "Verify your email and password. Contact an admin to check approval status."],
        ["Password reset not working", "Temporary password expired",
         "Ask an admin to reset your password from the Admin Panel."],
        ["Dashboard shows all zeros", "No data has been created yet",
         "Create your first strategy, add companies, and build a campaign to populate the dashboard."],
        ["AI Discovery returns no results", "ICP filters too restrictive",
         "Broaden your filters (e.g., add more industries or cities) and try again."],
        ["CSV import fails", "Incorrect column mapping or file format",
         "Ensure your CSV uses UTF-8 encoding and includes the required columns: Company Name, Industry, City."],
        ["Messages stuck in Draft", "Campaign not activated",
         "Ensure the campaign status is set to Active. Messages are generated only for active campaigns."],
        ["Zoho Mail not connecting", "Incorrect app password or SMTP settings",
         "Generate a new app-specific password in Zoho Accounts > Security > App Passwords."],
        ["Emails bouncing", "Invalid recipient email addresses",
         "Verify contact email addresses. Use the enrichment feature to validate emails."],
        ["Slow page loading", "Large dataset or network issues",
         "Use filters to reduce the displayed data. Check your internet connection."],
        ["Export download fails", "File too large or browser popup blocked",
         "Allow popups from the SalesPilot domain. Try exporting with a smaller date range."],
    ],
    col_widths=[1.6, 1.8, 3.2]
)

doc.add_heading('14.2  Getting Help', level=2)
doc.add_paragraph('If you encounter an issue not covered in this guide:')
add_bullet("Contact your SalesPilot administrator for account and access issues.")
add_bullet("Email support at support@clubconcierge.com for technical problems.")
add_bullet("Include screenshots and a description of the steps that led to the issue.")
add_bullet("Note the time, browser, and any error messages displayed.")

doc.add_heading('14.3  Keyboard Shortcuts', level=2)
add_styled_table(
    ["Shortcut", "Action"],
    [
        ["Ctrl + K  /  Cmd + K", "Open global search"],
        ["Ctrl + N  /  Cmd + N", "Create new record (context-dependent)"],
        ["Ctrl + S  /  Cmd + S", "Save current form"],
        ["Ctrl + E  /  Cmd + E", "Export current view"],
        ["Esc", "Close modal or cancel current action"],
        ["/", "Focus search bar on list pages"],
    ],
    col_widths=[2.2, 4.3]
)

doc.add_paragraph()
doc.add_paragraph()

# ─── Final page: Document info ──────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\u2014 End of Document \u2014")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.italic = True

doc.add_paragraph()

table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.rows[0].cells[0]
set_cell_shading(cell, "F8F9FA")
tc_pr = cell._tc.get_or_add_tcPr()
borders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tcBorders>'
)
tc_pr.append(borders)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SalesPilot Nigeria \u2014 User Manual v2.0")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = DARK_BLUE

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    "Published: April 2026\n"
    "Author: Club Concierge International\n"
    "Classification: Confidential\n"
    "Application: sales-management-nigeria.vercel.app"
)
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ─── Save ───────────────────────────────────────────────────────────────
output_path = "/Users/satguru/Projects/sales-nigeria/docs/SalesPilot Nigeria - User Manual v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
